from docx import Document

file_name = "origin_eng.txt"
wirte_file = "eng_6.txt"

words = []
count = 1

with open(file_name, 'r') as f:
    for line in f:
        if count % 3 is 2:
            # print line
            words.append(line)
        count = count + 1
f.close()

count = 1
tmp = ""
document = Document()

for word in words:
    word = word.strip('\n')
    tmp = tmp + word.ljust(13)
    if count % 6 is 0:
        print tmp
        # fw.write(tmp + '\n')
        document.add_paragraph(tmp)
        tmp = ""
    count = count + 1



# document.add_paragraph("hello word!")
document.save('demo.docx')